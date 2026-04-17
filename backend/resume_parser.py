# resume_parser.py
# Responsible for ONE thing only: extracting raw text from PDF or DOCX files.
# All analysis (skills, experience, questions) is done by the AI in question_generator.py

import PyPDF2
import docx
from io import BytesIO


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from every page of a PDF.
    Preserves line breaks between pages for better AI parsing.
    """
    text_parts = []
    try:
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
    except Exception as e:
        raise ValueError(f"Could not read PDF file: {str(e)}")

    if not text_parts:
        raise ValueError(
            "No text could be extracted from this PDF. "
            "Make sure it is not a scanned image — use a text-based PDF."
        )

    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file, including tables.
    Tables in resumes (like two-column layouts) are parsed row by row.
    """
    text_parts = []
    try:
        document = docx.Document(BytesIO(file_bytes))

        # Extract paragraph text
        for para in document.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract text from tables (many resumes use table layouts)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

    except Exception as e:
        raise ValueError(f"Could not read DOCX file: {str(e)}")

    if not text_parts:
        raise ValueError("No text could be extracted from this DOCX file.")

    return "\n".join(text_parts)


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Extract raw text from a resume file (PDF or DOCX).
    Returns a dict with the full resume text.
    All intelligence (skill extraction, question generation) happens via AI.
    """
    if not file_bytes:
        raise ValueError("The uploaded file is empty.")

    filename_lower = (filename or "").lower().strip()

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format '{filename}'. Please upload a PDF or DOCX file."
        )

    # Basic sanity check — a real resume should have at least 100 characters
    if len(text.strip()) < 100:
        raise ValueError(
            "The extracted text is too short. "
            "Make sure the file contains actual text content, not just images."
        )

    return {
        "text": text,
        "char_count": len(text),
        "filename": filename,
    }